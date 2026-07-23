#!/usr/bin/env python3
"""
Hilfsfunktionen für die DOCX-Report-Generierung bei Spezifikationsvergleichen.

Verwendung:
    from docx_helpers import ReportBuilder
    rb = ReportBuilder()
    rb.title_page("TR-01234", "Technische Richtlinie Beispielverfahren", "6.0", "7.0")
    rb.h1("1. Methodik")
    rb.p("Beschreibung...")
    rb.add_change_table([("Aspekt", "alt", "neu", "TECHNISCH")])
    rb.save("/pfad/zum/report.docx")
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import date


class ReportBuilder:
    """Erstellt einen formatierten Vergleichsreport als DOCX."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        for level, (sz, clr) in enumerate(
            [(16, '1F4E79'), (14, '2E75B6'), (12, '2E75B6')], 1
        ):
            hs = self.doc.styles[f'Heading {level}']
            hs.font.name = 'Arial'
            hs.font.size = Pt(sz)
            hs.font.color.rgb = RGBColor.from_string(clr)
            hs.font.bold = True

    def _setup_page(self):
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    # --- Titelseite ---

    def title_page(self, doc_id: str, doc_name: str, v_old: str, v_new: str):
        """Erstellt die Titelseite mit Bijektivitäts-Zusicherung."""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self._centered_run(f'Vollständige Vergleichsanalyse', 24, bold=True, color='1F4E79')
        self._centered_run(doc_id, 20, bold=True, color='1F4E79')
        self._centered_run(doc_name, 14, color='2E75B6')
        self._centered_run(f'Version {v_old} vs. Version {v_new}', 16, color='2E75B6')
        self.doc.add_paragraph()
        today = date.today().strftime('%d. %B %Y').replace(
            'January', 'Januar').replace('February', 'Februar').replace(
            'March', 'März').replace('May', 'Mai').replace(
            'June', 'Juni').replace('July', 'Juli').replace(
            'October', 'Oktober').replace('December', 'Dezember')
        self._centered_run(f'Erstellt am: {today}', 11, color='2E75B6')
        self._centered_run(
            'Methodik: Dreistufig (Maschineller Diff + Inhaltsanalyse + Kreuzvalidierung)',
            10, color='2E75B6'
        )
        self.doc.add_paragraph()
        p_note = self.doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_note.add_run(
            f'Dieses Dokument erfasst jede einzelne Änderung zwischen v{v_old} und v{v_new}.\n'
            f'Bijektive Vollständigkeit: v{v_old} + Änderungsdokument = v{v_new}.'
        )
        r.font.size = Pt(10)
        r.italic = True
        r.font.color.rgb = RGBColor.from_string('666666')
        self.page_break()

    def _centered_run(self, text, size, bold=False, color='000000'):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = RGBColor.from_string(color)
        r.font.name = 'Arial'
        return p

    # --- Überschriften ---

    def h1(self, text):
        return self.doc.add_heading(text, level=1)

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    # --- Absätze ---

    def p(self, text):
        return self.doc.add_paragraph(text)

    def pb(self, label: str, rest: str):
        """Absatz mit fettem Label gefolgt von normalem Text."""
        par = self.doc.add_paragraph()
        r1 = par.add_run(label)
        r1.bold = True
        par.add_run(rest)
        return par

    def bullet(self, text):
        return self.doc.add_paragraph(text, style='List Bullet')

    def bullet_bold(self, label: str, rest: str):
        """Aufzählungspunkt mit fettem Label."""
        par = self.doc.add_paragraph(style='List Bullet')
        r1 = par.add_run(label)
        r1.bold = True
        par.add_run(rest)
        return par

    # --- Tabellen ---

    def add_change_table(self, rows: list):
        """
        Erstellt eine Änderungstabelle.

        Args:
            rows: Liste von Tupeln (aspekt, v_alt, v_neu, kategorie)
        """
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        for i, t in enumerate(['Aspekt', 'v-alt', 'v-neu', 'Kategorie']):
            hdr[i].text = t
            for par in hdr[i].paragraphs:
                for run in par.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
        for aspekt, v_alt, v_neu, kat in rows:
            row = table.add_row().cells
            row[0].text = str(aspekt)
            row[1].text = str(v_alt)
            row[2].text = str(v_neu)
            row[3].text = str(kat)
            for c in row:
                for par in c.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'
        for row in table.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(5.5)
            row.cells[2].width = Cm(5.5)
            row.cells[3].width = Cm(2)
        return table

    def add_info_table(self, headers: list, rows: list):
        """
        Erstellt eine allgemeine Informationstabelle.

        Args:
            headers: Liste von Spaltenüberschriften
            rows: Liste von Tupeln mit Zellwerten
        """
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, t in enumerate(headers):
            hdr[i].text = t
            for par in hdr[i].paragraphs:
                for run in par.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for r in rows:
            row = table.add_row().cells
            for i, v in enumerate(r):
                row[i].text = str(v)
                for par in row[i].paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
        return table

    # --- Seitenumbruch ---

    def page_break(self):
        self.doc.add_page_break()

    # --- Speichern ---

    def save(self, path: str):
        os.makedirs(os.path.dirname(path) or '.', exist_ok=True)
        self.doc.save(path)
        size = os.path.getsize(path)
        print(f"DOCX saved: {path} ({size:,} bytes)")
        return path
